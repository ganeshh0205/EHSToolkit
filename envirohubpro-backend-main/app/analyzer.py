import logging
from pathlib import Path
from typing import Any
import pandas as pd

logger = logging.getLogger("envirohubpro.analyzer")


class DataAnalyzer:
    def __init__(self) -> None:
        self.data: pd.DataFrame | None = None

    def load_data(self, file_path: str | Path) -> pd.DataFrame:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        suffix = path.suffix.lower()
        try:
            if suffix == ".csv":
                self.data = pd.read_csv(path)
            elif suffix in [".xlsx", ".xls"]:
                self.data = pd.read_excel(path)
            elif suffix == ".json":
                self.data = pd.read_json(path)
            elif suffix == ".pdf":
                try:
                    import pdfplumber
                except ImportError:
                    raise ImportError("pdfplumber is required to parse PDFs")
                
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        table = page.extract_table()
                        if table:
                            # Use first found table
                            headers = table[0]
                            self.data = pd.DataFrame(table[1:], columns=headers)
                            break
                    if self.data is None:
                        raise ValueError("No extractable tabular data found in the PDF.")
            elif suffix == ".docx":
                try:
                    from docx import Document
                except ImportError:
                    raise ImportError("python-docx is required to parse DOCX")
                
                doc = Document(path)
                if not doc.tables:
                    raise ValueError("No tables found in the DOCX file.")
                
                table = doc.tables[0]
                data = []
                headers = []
                for i, row in enumerate(table.rows):
                    text = [cell.text.strip() for cell in row.cells]
                    if i == 0:
                        headers = text
                    else:
                        data.append(text)
                
                self.data = pd.DataFrame(data, columns=headers)
            else:
                raise ValueError(f"Unsupported file format: {suffix}")

            logger.info(f"Loaded {len(self.data)} rows from {path.name}")
            return self.data
        except Exception as e:
            logger.error(f"Error loading {path.name}: {e}")
            raise

    def clean_data(self) -> pd.DataFrame:
        if self.data is None:
            raise ValueError("No data loaded. Call load_data() first.")

        self.data.drop_duplicates(inplace=True)
        self.data.dropna(how="all", inplace=True)

        for col in self.data.select_dtypes(include=["object", "string"]):
            self.data[col] = self.data[col].apply(lambda x: x.strip() if isinstance(x, str) else x)

        return self.data

    def transform_data(self, date_columns: list[str] | None = None) -> pd.DataFrame:
        if self.data is None:
            raise ValueError("No data to transform.")

        self.data.columns = [
            col.strip().lower().replace(" ", "_").replace("-", "_")
            for col in self.data.columns
        ]

        if date_columns:
            for col in date_columns:
                if col in self.data.columns:
                    self.data[col] = pd.to_datetime(self.data[col], errors="coerce")

        # Check if data needs reshaping from wide to long
        analyte_synonyms = ['analyte', 'parameter', 'chemical', 'pollutant', 'constituent', 'substance', 'material', 'indicator']
        # Do not treat 'material' as analyte if it's strictly an ID column like in Asbestos dataset
        # Actually 'parameter' and 'analyte' are the most common
        has_analyte_col = any(any(syn in col for syn in analyte_synonyms) for col in self.data.columns)
        
        # Exception: if there is a 'material' column but also chemical columns like 'asbestos_percent'
        # 'material' is usually an ID var in this context.
        if has_analyte_col and 'material' in self.data.columns and any('percent' in c or 'mgkg' in c for c in self.data.columns):
            has_analyte_col = False
            
        if not has_analyte_col:
            # Likely wide data format. Attempt to melt.
            # Identify potential ID columns (date, id, location, etc.)
            id_keywords = ['id', 'date', 'location', 'depth', 'facility', 'stack', 'station', 'material', 'well', 'time', 'site']
            id_vars = [col for col in self.data.columns if any(kw in col for kw in id_keywords)]
            value_vars = [col for col in self.data.columns if col not in id_vars]
            
            if value_vars:
                melted = self.data.melt(id_vars=id_vars, value_vars=value_vars, var_name='original_column', value_name='result')
                
                def extract_unit(name: str):
                    if not isinstance(name, str):
                        return name, ""
                    parts = name.rsplit('_', 1)
                    if len(parts) == 2:
                        unit_str = parts[1]
                        unit_str = unit_str.replace('mgkg', 'mg/kg').replace('ugm3', 'ug/m3').replace('mgl', 'mg/L')
                        return parts[0], unit_str
                    return name, ""
                
                extracted = melted['original_column'].apply(lambda x: pd.Series(extract_unit(x)))
                melted['analyte'] = extracted[0]
                melted['unit'] = extracted[1]
                
                melted.dropna(subset=['result'], inplace=True)
                self.data = melted

        return self.data

    def run_environmental_analysis(self) -> dict[str, Any]:
        if self.data is None:
            raise ValueError("Data must be loaded before analysis.")

        return {
            "summary": {
                "rows": len(self.data),
                "columns": list(self.data.columns),
            },
            "status": "In-progress",
            "message": "Environmental analysis logic to be implemented."
        }
